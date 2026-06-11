from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path

from fastapi import HTTPException, UploadFile, status
from pydantic import ValidationError

from app.schemas.onboarding import (
    ExtractedCandidateProfile,
    ExtractedProject,
    ExtractedWorkExperience,
    ResumeConfidence,
    ResumeParseDraft,
    ResumeParseResponse,
)
from app.services.ai_provider import ProviderOutputError
from app.services.ai_provider_factory import build_ai_provider

MAX_RESUME_BYTES = 5 * 1024 * 1024
MAX_AI_TEXT_CHARS = 12000
MIN_USEFUL_TEXT_CHARS = 100
ALLOWED_CONTENT_TYPES = {
    "application/pdf": ".pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
}
ALLOWED_EXTENSIONS = {".pdf", ".docx"}
IMPORTANT_MISSING_FIELDS = (
    "full_name",
    "target_role",
    "university",
    "degree",
    "skills",
    "tech_stack",
    "projects",
    "gpa",
    "github_url/portfolio_url",
)
POLLUTED_FIELD_WARNING = "Some resume fields were uncertain and were left blank for manual review."
DEGREE_TERMS = (
    "bs",
    "b.s",
    "bsc",
    "b.sc",
    "bachelor",
    "ms",
    "m.s",
    "msc",
    "m.sc",
    "master",
    "computer science",
    "software engineering",
    "artificial intelligence",
    " ai",
    "data science",
)
KNOWN_SKILL_TERMS = (
    "react",
    "next.js",
    "typescript",
    "javascript",
    "fastapi",
    "python",
    "postgresql",
    "sql",
    "node",
    "express",
    "docker",
    "tailwind",
    "machine learning",
    "mongodb",
)
EMAIL_PATTERN = re.compile(r"[\w.+-]+@[\w-]+(?:\.[\w-]+)+")
PHONE_PATTERN = re.compile(r"(?<!\d)(?:\+?\d[\d\s().-]{7,}\d)(?!\d)")
URL_PATTERN = re.compile(r"(?:https?://|www\.)[^\s,;)]+|(?:github\.com|linkedin\.com)/[^\s,;)]+", re.IGNORECASE)
SECTION_HEADERS = (
    "professional summary",
    "core skills",
    "skills",
    "relevant experience",
    "work experience",
    "selected projects",
    "projects",
    "education and certifications",
    "education",
    "certifications",
)


def _normalize_whitespace(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip()


def _validate_upload_metadata(file: UploadFile) -> str:
    extension = Path(file.filename or "").suffix.lower()
    content_type = (file.content_type or "").split(";")[0].strip().lower()
    if extension not in ALLOWED_EXTENSIONS or content_type not in ALLOWED_CONTENT_TYPES:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
            detail={
                "detail": "Unsupported resume file type. Upload a PDF or DOCX file.",
                "reason": "unsupported_resume_type",
            },
        )
    expected_extension = ALLOWED_CONTENT_TYPES[content_type]
    if extension != expected_extension:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
            detail={
                "detail": "Resume file extension does not match its content type.",
                "reason": "unsupported_resume_type",
            },
        )
    return extension


async def _read_upload_bytes(file: UploadFile) -> bytes:
    data = await file.read()
    if len(data) > MAX_RESUME_BYTES:
        raise HTTPException(
            status_code=status.HTTP_413_CONTENT_TOO_LARGE,
            detail={
                "detail": "Resume file is too large. Upload a file up to 5MB.",
                "reason": "resume_file_too_large",
            },
        )
    return data


def extract_pdf_text(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ProviderOutputError("pypdf dependency is not installed") from exc

    reader = PdfReader(BytesIO(data))
    parts = [(page.extract_text() or "") for page in reader.pages]
    return _normalize_whitespace("\n".join(parts))


def extract_docx_text(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise ProviderOutputError("python-docx dependency is not installed") from exc

    document = Document(BytesIO(data))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return _normalize_whitespace("\n".join(parts))


def extract_resume_text(data: bytes, extension: str) -> str:
    if extension == ".pdf":
        return extract_pdf_text(data)
    if extension == ".docx":
        return extract_docx_text(data)
    return ""


def ensure_useful_resume_text(text: str) -> str:
    normalized = _normalize_whitespace(text)
    useful_chars = len(re.sub(r"[^A-Za-z0-9]", "", normalized))
    if useful_chars < MIN_USEFUL_TEXT_CHARS:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
            detail={
                "detail": "Resume text could not be read. Please enter details manually.",
                "reason": "resume_text_unreadable",
            },
        )
    return normalized


def extract_email_from_text(text: str) -> str | None:
    match = EMAIL_PATTERN.search(text)
    return match.group(0) if match else None


def extract_phone_from_text(text: str) -> str | None:
    for match in PHONE_PATTERN.finditer(text):
        value = match.group(0).strip()
        digits = re.sub(r"\D", "", value)
        if 8 <= len(digits) <= 15:
            return value
    return None


def _normalize_url(value: str) -> str:
    cleaned = value.strip().rstrip(".,;)")
    if cleaned.startswith("www."):
        return f"https://{cleaned}"
    if cleaned.lower().startswith(("github.com/", "linkedin.com/")):
        return f"https://{cleaned}"
    return cleaned


def extract_urls_from_text(text: str) -> list[str]:
    urls: list[str] = []
    seen: set[str] = set()
    for match in URL_PATTERN.finditer(text):
        url = _normalize_url(match.group(0))
        key = url.lower()
        if key not in seen:
            urls.append(url)
            seen.add(key)
    return urls


def _text_lines(text: str) -> list[str]:
    return [line.strip() for line in re.split(r"[\r\n]+", text) if line.strip()]


def _compact_text(text: str) -> str:
    return _normalize_whitespace(text)


def _looks_like_name(value: str) -> bool:
    cleaned = value.strip()
    if not cleaned or re.search(r"\d|@|https?://|www\.|github\.com|linkedin\.com", cleaned, flags=re.IGNORECASE):
        return False
    words = cleaned.split()
    if not 2 <= len(words) <= 5:
        return False
    return all(re.fullmatch(r"[A-Za-z][A-Za-z.'-]*", word) for word in words)


def _extract_header_name_and_role(text: str) -> tuple[str | None, str | None]:
    lines = _text_lines(text)
    if len(lines) >= 2 and _looks_like_name(lines[0]):
        title = clean_scalar(lines[1], max_words=8, max_chars=80)
        return lines[0].strip(), title

    compact = _compact_text(text)
    name_match = re.match(r"^([A-Z][A-Z.'-]+(?:\s+[A-Z][A-Z.'-]+){1,4})\s+", compact)
    if not name_match:
        return None, None
    name = name_match.group(1).strip()
    remainder = compact[name_match.end() :].strip()
    title = None
    role_match = re.match(
        r"(.{2,80}?\b(?:engineer|developer|scientist|analyst|designer|manager|specialist|architect|consultant|intern|administrator))\b\s+"
        r"(?:[A-Z][A-Za-z .'-]+,\s*[A-Z][A-Za-z .'-]+\s*\||\+?\d|[\w.+-]+@|linkedin\.com|github\.com|professional summary)",
        remainder,
        flags=re.IGNORECASE,
    )
    if role_match:
        title = role_match.group(1).strip()
        return name, clean_scalar(title, max_words=8, max_chars=80)

    location_match = re.match(r"(.+?)\s+[A-Z][A-Za-z .'-]+,\s*[A-Z][A-Za-z .'-]+\s*\|", remainder)
    if location_match:
        title = location_match.group(1).strip()
    else:
        contact_match = re.match(r"(.+?)\s+(?:\+?\d|[\w.+-]+@|linkedin\.com|github\.com|professional summary)", remainder, flags=re.IGNORECASE)
        if contact_match:
            title = contact_match.group(1).strip()
    return name, clean_scalar(title, max_words=8, max_chars=80)


def _section_text(text: str, section_names: tuple[str, ...]) -> str:
    compact = _compact_text(text)
    lowered = compact.lower()
    matches = [(lowered.find(name), name) for name in section_names if lowered.find(name) != -1]
    if not matches:
        return ""
    start, matched_name = min(matches, key=lambda item: item[0])
    content_start = start + len(matched_name)
    end = len(compact)
    for header in SECTION_HEADERS:
        header_index = lowered.find(header, content_start)
        if header_index != -1 and header_index > content_start:
            end = min(end, header_index)
    return compact[start:end].strip()


def _extract_education_basics(text: str) -> dict:
    education = _section_text(text, ("education and certifications", "education"))
    result = {"university": None, "degree": None, "graduation_year": None}
    if not education:
        return result

    education = re.sub(r"(?i)^education(?:\s+and\s+certifications)?\s*", "", education).strip()
    education = re.split(r"(?i)\bcertifications?\s*:", education, maxsplit=1)[0].strip()
    edu_line_match = re.search(
        r"(?P<degree>(?:BS|BSc|B\.S\.?|Bachelor|MS|MSc|M\.S\.?)[^|–—\n-]{0,80})\s*[-–—]\s*"
        r"(?P<university>[^|\n]{2,120})(?:\|\s*(?P<year_text>[^|\n]{0,80}))?",
        education,
        flags=re.IGNORECASE,
    )
    if edu_line_match:
        result["degree"] = clean_scalar(edu_line_match.group("degree"), max_words=12, max_chars=120)
        result["university"] = clean_scalar(edu_line_match.group("university"), max_words=12, max_chars=120)
        year_text = edu_line_match.group("year_text") or edu_line_match.group(0)
    else:
        year_text = education

    year_patterns = (
        r"(?i)\b(?:expected|graduating|graduation|class of)\s*(20\d{2}|19\d{2})\b",
        r"(?i)\b(?:degree|university|computer science|software engineering|data science|ai)\b[^.|\n]{0,80}\b(20\d{2}|19\d{2})\b",
    )
    for pattern in year_patterns:
        match = re.search(pattern, year_text)
        if match:
            result["graduation_year"] = int(match.group(1))
            break
    return result


def _extract_explicit_gpa(text: str) -> float | None:
    match = re.search(r"(?i)\b(?:gpa|cgpa|grade point average)\s*[:\-]?\s*(\d+(?:\.\d+)?)\b", text)
    if not match:
        return None
    return _normalize_gpa(match.group(1))


def extract_resume_basics_deterministic(resume_text: str) -> dict:
    name, title = _extract_header_name_and_role(resume_text)
    urls = extract_urls_from_text(resume_text)
    education = _extract_education_basics(resume_text)
    return {
        "full_name": name,
        "target_role": title,
        "email": extract_email_from_text(resume_text),
        "phone": extract_phone_from_text(resume_text),
        "linkedin_url": _first_url(urls, "linkedin"),
        "github_url": _first_url(urls, "github"),
        "university": education["university"],
        "degree": education["degree"],
        "graduation_year": education["graduation_year"],
        "gpa": _extract_explicit_gpa(resume_text),
    }


def _contains_email(value: str) -> bool:
    return bool(EMAIL_PATTERN.search(value))


def _contains_phone(value: str) -> bool:
    return extract_phone_from_text(value) is not None


def _urls_in_value(value: str) -> list[str]:
    return extract_urls_from_text(value)


def _word_count(value: str) -> int:
    return len(re.findall(r"[A-Za-z0-9+#.]+", value))


def _skill_term_count(value: str) -> int:
    lowered = value.lower()
    return sum(1 for term in KNOWN_SKILL_TERMS if term in lowered)


def looks_like_resume_dump(value) -> bool:
    if value is None:
        return False
    text = str(value).strip()
    if not text:
        return False
    lowered = text.lower()
    if "\n" in text or "\r" in text or any(marker in text for marker in ["•", "●", "▪", " - "]):
        return True
    urls = _urls_in_value(text)
    if len(urls) > 1:
        return True
    if _contains_email(text) and (_contains_phone(text) or urls or _skill_term_count(text) >= 1):
        return True
    if _contains_phone(text) and urls:
        return True
    if len(text) > 220 or _word_count(text) > 28:
        return True
    section_markers = ("skills", "projects", "experience", "education", "summary", "certifications")
    if sum(1 for marker in section_markers if marker in lowered) >= 2:
        return True
    if _skill_term_count(text) >= 3:
        return True
    if text.count(",") >= 4 or text.count("|") >= 3:
        return True
    return False


def clean_scalar(value, max_words: int | None = None, max_chars: int | None = None) -> str | None:
    if value is None:
        return None
    text = str(value).strip()
    if not text or looks_like_resume_dump(text):
        return None
    text = _normalize_whitespace(text)
    if max_chars is not None and len(text) > max_chars:
        return None
    if max_words is not None and _word_count(text) > max_words:
        return None
    return text


def _clean_non_contact_scalar(value, *, max_words: int | None = None, max_chars: int | None = None) -> str | None:
    text = clean_scalar(value, max_words=max_words, max_chars=max_chars)
    if not text or _contains_email(text) or _contains_phone(text) or _urls_in_value(text):
        return None
    return text


def _looks_degree_like(value: str | None) -> bool:
    if not value:
        return False
    lowered = f" {value.lower()} "
    return any(term in lowered for term in DEGREE_TERMS)


def _clean_url(value: str | None, *, kind: str | None = None) -> str | None:
    if not value:
        return None
    raw = str(value).strip().rstrip(".,;)")
    urls = _urls_in_value(raw)
    if len(urls) != 1:
        return None
    url = urls[0]
    if _normalize_url(raw) != url:
        return None
    lowered = url.lower()
    if kind == "github" and "github.com" not in lowered:
        return None
    if kind == "linkedin" and "linkedin.com" not in lowered:
        return None
    if kind == "portfolio" and ("github.com" in lowered or "linkedin.com" in lowered):
        return None
    return url


def _first_url(urls: list[str], kind: str) -> str | None:
    for url in urls:
        lowered = url.lower()
        if kind == "github" and "github.com" in lowered:
            return url
        if kind == "linkedin" and "linkedin.com" in lowered:
            return url
        if kind == "portfolio" and "github.com" not in lowered and "linkedin.com" not in lowered:
            return url
    return None


def _split_skill_value(value: str) -> list[str]:
    return [item.strip() for item in re.split(r"[,;/|]", value) if item.strip()]


def _clean_string_list(values: list[str]) -> list[str]:
    cleaned: list[str] = []
    seen: set[str] = set()
    for value in values or []:
        candidates = _split_skill_value(str(value)) if any(separator in str(value) for separator in [",", ";", "/", "|"]) else [str(value)]
        for candidate in candidates:
            text = _clean_non_contact_scalar(candidate, max_words=5, max_chars=60)
            if not text:
                continue
            key = text.lower()
            if key not in seen:
                cleaned.append(text)
                seen.add(key)
    return cleaned


def _extract_name_from_text(text: str) -> str | None:
    lines = [line.strip() for line in re.split(r"[\r\n]+", text) if line.strip()]
    for line in lines[:8]:
        candidate = _clean_non_contact_scalar(line, max_words=5, max_chars=120)
        if not candidate or any(marker in candidate.lower() for marker in ["resume", "curriculum", "skills", "education"]):
            continue
        if re.search(r"\d", candidate):
            continue
        if 2 <= _word_count(candidate) <= 5:
            return candidate
    email = extract_email_from_text(text)
    if email:
        prefix = text[: text.find(email)].strip(" \t\r\n-|,")
        words = prefix.split()[-5:]
        candidate = " ".join(words)
        if candidate and 2 <= _word_count(candidate) <= 5 and not re.search(r"\d", candidate):
            return candidate
    return None


def _clean_optional_text(value: str | None) -> str | None:
    if value is None:
        return None
    trimmed = str(value).strip()
    return trimmed or None


def _normalize_gpa(value: float | int | str | None) -> float | None:
    if value in (None, ""):
        return None
    try:
        numeric = float(value)
    except (TypeError, ValueError):
        return None
    if numeric < 0:
        return None
    if numeric <= 10 or numeric <= 100:
        return numeric
    return None


def _normalize_profile(profile: ExtractedCandidateProfile) -> ExtractedCandidateProfile:
    data = profile.model_dump()
    for field_name, value in list(data.items()):
        if isinstance(value, str) or value is None:
            data[field_name] = _clean_optional_text(value)
    data["gpa"] = _normalize_gpa(data.get("gpa"))
    return ExtractedCandidateProfile.model_validate(data)


def sanitize_extracted_profile(profile: ExtractedCandidateProfile, resume_text: str) -> tuple[ExtractedCandidateProfile, bool]:
    polluted = False
    urls = extract_urls_from_text(resume_text)

    full_name = _clean_non_contact_scalar(profile.full_name, max_words=5, max_chars=120)
    if full_name != profile.full_name:
        polluted = bool(profile.full_name)
    if not full_name:
        full_name = _extract_name_from_text(resume_text)

    university = _clean_non_contact_scalar(profile.university, max_words=10, max_chars=120)
    if university and _skill_term_count(university) >= 2:
        university = None
    if university != profile.university:
        polluted = polluted or bool(profile.university)

    degree = _clean_non_contact_scalar(profile.degree, max_words=12, max_chars=120)
    if degree and not _looks_degree_like(degree):
        degree = None
    if degree != profile.degree:
        polluted = polluted or bool(profile.degree)

    target_role = _clean_non_contact_scalar(profile.target_role, max_words=8, max_chars=80)
    if target_role and any(marker in target_role.lower() for marker in ["skills", "education", "experience", "project"]):
        target_role = None
    if target_role != profile.target_role:
        polluted = polluted or bool(profile.target_role)

    email = extract_email_from_text(resume_text) or clean_scalar(profile.email, max_words=1, max_chars=254)
    phone = extract_phone_from_text(resume_text) or clean_scalar(profile.phone, max_words=5, max_chars=80)
    github_url = _first_url(urls, "github") or _clean_url(profile.github_url, kind="github")
    linkedin_url = _first_url(urls, "linkedin") or _clean_url(profile.linkedin_url, kind="linkedin")
    portfolio_url = _first_url(urls, "portfolio") or _clean_url(profile.portfolio_url, kind="portfolio")

    skills = _clean_string_list(profile.skills)
    tech_stack = _clean_string_list(profile.tech_stack)

    projects: list[ExtractedProject] = []
    for project in profile.projects:
        title = _clean_non_contact_scalar(project.title, max_words=14, max_chars=160)
        description = clean_scalar(project.description, max_words=90, max_chars=900)
        technologies = _clean_string_list(project.technologies)
        github_project_url = _clean_url(project.github_url, kind="github")
        live_url = _clean_url(project.live_url)
        if title or description:
            projects.append(
                ExtractedProject(
                    title=title,
                    description=description,
                    technologies=technologies,
                    github_url=github_project_url,
                    live_url=live_url,
                )
            )

    work_experience: list[ExtractedWorkExperience] = []
    for item in profile.work_experience:
        company = _clean_non_contact_scalar(item.company, max_words=8, max_chars=120)
        role = _clean_non_contact_scalar(item.role, max_words=8, max_chars=120)
        duration = clean_scalar(item.duration, max_words=8, max_chars=120)
        description = clean_scalar(item.description, max_words=80, max_chars=800)
        if company or role or description:
            work_experience.append(
                ExtractedWorkExperience(
                    company=company,
                    role=role,
                    duration=duration,
                    description=description,
                )
            )

    sanitized = ExtractedCandidateProfile(
        full_name=full_name,
        email=email,
        phone=phone,
        university=university,
        degree=degree,
        graduation_year=profile.graduation_year,
        gpa=profile.gpa,
        target_role=target_role,
        experience_level=profile.experience_level,
        skills=skills,
        tech_stack=tech_stack,
        projects=projects,
        work_experience=work_experience,
        github_url=github_url,
        linkedin_url=linkedin_url,
        portfolio_url=portfolio_url,
    )
    return sanitized, polluted


def _merge_deterministic_profile(profile: ExtractedCandidateProfile, basics: dict) -> ExtractedCandidateProfile:
    data = profile.model_dump()
    for field_name in [
        "full_name",
        "email",
        "phone",
        "linkedin_url",
        "github_url",
        "university",
        "degree",
        "target_role",
    ]:
        value = basics.get(field_name)
        if value not in (None, "", []):
            data[field_name] = value

    # These fields are scoped deterministically. If absent in their allowed section,
    # keep them null instead of accepting AI guesses from work/project dates or defaults.
    data["graduation_year"] = basics.get("graduation_year")
    data["gpa"] = basics.get("gpa")
    return ExtractedCandidateProfile.model_validate(data)


def _warnings_for_missing_fields(profile: ExtractedCandidateProfile, warnings: list[str]) -> list[str]:
    recomputed_missing_warnings = {
        "full_name missing or uncertain",
        "target_role missing or uncertain",
        "university missing or uncertain",
        "degree missing or uncertain",
        "skills missing",
        "tech_stack missing",
        "projects missing",
        "gpa missing",
        "github_url/portfolio_url missing",
    }
    next_warnings = [
        warning.strip()
        for warning in warnings
        if warning and warning.strip() and warning.strip() not in recomputed_missing_warnings
    ]
    next_warnings = list(dict.fromkeys(next_warnings))
    if not profile.full_name:
        next_warnings.append("full_name missing or uncertain")
    if not profile.target_role:
        next_warnings.append("target_role missing or uncertain")
    if not profile.university:
        next_warnings.append("university missing or uncertain")
    if not profile.degree:
        next_warnings.append("degree missing or uncertain")
    if not profile.skills:
        next_warnings.append("skills missing")
    if not profile.tech_stack:
        next_warnings.append("tech_stack missing")
    if not profile.projects:
        next_warnings.append("projects missing")
    if profile.gpa is None:
        next_warnings.append("gpa missing")
    if not profile.github_url and not profile.portfolio_url:
        next_warnings.append("github_url/portfolio_url missing")
    return list(dict.fromkeys(next_warnings))


def parse_resume_text_with_ai(resume_text: str, provider_name: str | None = None) -> tuple[ResumeParseDraft, object]:
    provider = build_ai_provider(provider_name, capability="onboarding")
    basics = extract_resume_basics_deterministic(resume_text)
    try:
        draft = provider.parse_resume_profile(resume_text[:MAX_AI_TEXT_CHARS])
    except Exception as exc:
        raise ProviderOutputError("Resume parsing failed") from exc
    try:
        profile = _normalize_profile(draft.extracted_profile)
        profile = _merge_deterministic_profile(profile, basics)
        profile, polluted = sanitize_extracted_profile(profile, resume_text)
        confidence = ResumeConfidence.model_validate(draft.confidence.model_dump())
        warnings = list(draft.warnings)
        if polluted:
            warnings.append(POLLUTED_FIELD_WARNING)
        return (
            ResumeParseDraft(
                extracted_profile=profile,
                confidence=confidence,
                warnings=_warnings_for_missing_fields(profile, warnings),
            ),
            provider,
        )
    except (ValidationError, ValueError, TypeError) as exc:
        raise ProviderOutputError("Resume parsing returned invalid profile data") from exc


async def parse_resume_upload(file: UploadFile, provider_name: str | None = None) -> ResumeParseResponse:
    extension = _validate_upload_metadata(file)
    data = await _read_upload_bytes(file)
    text = ensure_useful_resume_text(extract_resume_text(data, extension))
    draft, provider = parse_resume_text_with_ai(text, provider_name=provider_name)
    return ResumeParseResponse(
        status="parsed",
        extracted_profile=draft.extracted_profile,
        confidence=draft.confidence,
        warnings=draft.warnings,
        raw_text_preview=text[:500],
        provider_metadata=provider.state.metadata(),
    )
